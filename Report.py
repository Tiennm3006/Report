import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from datetime import datetime
from docx import Document
from io import BytesIO
import openpyxl
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference, Series
from openpyxl.utils.dataframe import dataframe_to_rows

st.set_page_config(page_title="Phân tích Điện lực", layout="wide")
st.title("🔍 Báo cáo phân tích kết quả kiểm tra hệ thống điện")

uploaded_file = st.file_uploader("Tải lên file Excel chứa sheet 'Tong hop luy ke'", type=["xlsx"])

if uploaded_file:
    df = pd.read_excel(uploaded_file, sheet_name="Tong hop luy ke", header=None)

    # Làm sạch dữ liệu
    df_cleaned = df.iloc[4:].copy()
    df_cleaned.columns = [
        "STT", "Điện lực", "1P_GT", "1P_TT", "3P_GT", "3P_TT",
        "TU", "TI", "Tổng công tơ", "Kế hoạch", "Tỷ lệ"
    ]
    df_cleaned = df_cleaned[df_cleaned["Điện lực"].notna()]
    cols_to_numeric = ["1P_GT", "1P_TT", "3P_GT", "3P_TT", "TU", "TI", "Tổng công tơ", "Kế hoạch", "Tỷ lệ"]
    df_cleaned[cols_to_numeric] = df_cleaned[cols_to_numeric].apply(pd.to_numeric, errors='coerce')
    df_cleaned["Tỷ lệ"] = df_cleaned["Tỷ lệ"] * 100  # đổi sang %

    # Tổng quan và dự báo
    total_current = df_cleaned["Tổng công tơ"].sum()
    total_plan = df_cleaned["Kế hoạch"].sum()
    current_date = datetime.now()
    days_passed = (current_date - datetime(2025, 1, 1)).days
    days_total = (datetime(2025, 9, 30) - datetime(2025, 1, 1)).days
    avg_per_day = total_current / days_passed
    forecast_total = avg_per_day * days_total
    forecast_ratio = forecast_total / total_plan

    # Top 3 và Bottom 3
    df_sorted = df_cleaned.sort_values(by="Tỷ lệ", ascending=False)
    top_3 = df_sorted.head(3)
    bottom_3 = df_sorted.tail(3)

    # Hiển thị số liệu
    st.subheader("Tổng quan và dự báo")
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Tổng đã thực hiện", f"{total_current:,}")
        st.metric("Kế hoạch", f"{total_plan:,}")
        st.metric("Tốc độ TB/ngày", f"{avg_per_day:.2f}")
    with col2:
        st.metric("Dự báo đến 30/09/2025", f"{int(forecast_total):,}")
        st.metric("Tỷ lệ dự báo", f"{forecast_ratio*100:.2f}%")

    st.subheader("Biểu đồ tỷ lệ hoàn thành của các Điện lực")
    fig, ax = plt.subplots(figsize=(10, 5))
    bars = ax.bar(df_sorted["Điện lực"], df_sorted["Tỷ lệ"])
    ax.set_ylabel("Tỷ lệ hoàn thành (%)")
    ax.set_title("Tỷ lệ hoàn thành kế hoạch theo Điện lực")
    ax.tick_params(axis='x', rotation=45)
    for bar in bars:
        height = bar.get_height()
        ax.annotate(f"{height:.1f}%", xy=(bar.get_x() + bar.get_width() / 2, height),
                    xytext=(0, 3), textcoords="offset points", ha='center', va='bottom')
    st.pyplot(fig)

    st.subheader("Biểu đồ Top 3 tỷ lệ cao nhất")
    fig_top, ax_top = plt.subplots()
    bars_top = ax_top.bar(top_3["Điện lực"], top_3["Tỷ lệ"], color='green')
    ax_top.set_ylabel("Tỷ lệ hoàn thành (%)")
    ax_top.set_title("Top 3 Điện lực có tỷ lệ hoàn thành cao nhất")
    for bar in bars_top:
        height = bar.get_height()
        ax_top.annotate(f"{height:.1f}%", xy=(bar.get_x() + bar.get_width() / 2, height),
                        xytext=(0, 3), textcoords="offset points", ha='center', va='bottom')
    st.pyplot(fig_top)

    st.subheader("Biểu đồ Bottom 3 tỷ lệ thấp nhất")
    fig_bot, ax_bot = plt.subplots()
    bars_bot = ax_bot.bar(bottom_3["Điện lực"], bottom_3["Tỷ lệ"], color='red')
    ax_bot.set_ylabel("Tỷ lệ hoàn thành (%)")
    ax_bot.set_title("Bottom 3 Điện lực có tỷ lệ hoàn thành thấp nhất")
    for bar in bars_bot:
        height = bar.get_height()
        ax_bot.annotate(f"{height:.1f}%", xy=(bar.get_x() + bar.get_width() / 2, height),
                        xytext=(0, 3), textcoords="offset points", ha='center', va='bottom')
    st.pyplot(fig_bot)

    st.subheader("Top 3 tỷ lệ cao nhất")
    st.dataframe(top_3[["Điện lực", "Tổng công tơ", "Kế hoạch", "Tỷ lệ"]])

    st.subheader("Bottom 3 tỷ lệ thấp nhất")
    st.dataframe(bottom_3[["Điện lực", "Tổng công tơ", "Kế hoạch", "Tỷ lệ"]])

    # Tạo file Excel báo cáo
    def generate_excel():
        wb = Workbook()
        ws = wb.active
        ws.title = "Tong hop"

        for r in dataframe_to_rows(df_sorted, index=False, header=True):
            ws.append(r)

        def add_chart(sheet, title, data_col, label_col, start_cell):
            chart = BarChart()
            chart.title = title
            chart.y_axis.title = "Tỷ lệ (%)"
            data = Reference(sheet, min_col=data_col, min_row=1, max_row=sheet.max_row)
            cats = Reference(sheet, min_col=label_col, min_row=2, max_row=sheet.max_row)
            chart.add_data(data, titles_from_data=True)
            chart.set_categories(cats)
            sheet.add_chart(chart, start_cell)

        add_chart(ws, "Tỷ lệ hoàn thành theo Điện lực", data_col=11, label_col=2, start_cell="M5")

        # Top 3 sheet
        ws_top = wb.create_sheet("Top 3")
        for r in dataframe_to_rows(top_3, index=False, header=True):
            ws_top.append(r)
        add_chart(ws_top, "Top 3 tỷ lệ cao nhất", data_col=11, label_col=2, start_cell="M5")

        # Bottom 3 sheet
        ws_bot = wb.create_sheet("Bottom 3")
        for r in dataframe_to_rows(bottom_3, index=False, header=True):
            ws_bot.append(r)
        add_chart(ws_bot, "Bottom 3 tỷ lệ thấp nhất", data_col=11, label_col=2, start_cell="M5")

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_docx():
        doc = Document()
        doc.add_heading('BÁO CÁO PHÂN TÍCH KẾT QUẢ KIỂM TRA HỆ THỐNG ĐẾN NGÀY ' + current_date.strftime('%d/%m/%Y'), 0)
        doc.add_heading('I. ĐÁNH GIÁ TỔNG QUÁT', level=1)
        doc.add_paragraph(
            f"Tổng số công tơ đã thực hiện: {total_current:,}\n"
            f"Kế hoạch tổng: {total_plan:,}\n"
            f"Số ngày đã thực hiện: {days_passed} ngày\n"
            f"Tốc độ trung bình/ngày: {avg_per_day:.2f} công tơ/ngày\n"
            f"Dự báo đến 30/09/2025: {int(forecast_total):,} công tơ\n"
            f"Tỷ lệ dự báo: {forecast_ratio*100:.2f}%\n"
        )
        doc.add_heading('II. TOP 3 ĐIỆN LỰC', level=1)
        for _, row in top_3.iterrows():
            doc.add_paragraph(
                f"- {row['Điện lực']}: {row['Tổng công tơ']:,}/{row['Kế hoạch']:,} ({row['Tỷ lệ']:.2f}%)"
            )
        doc.add_heading('III. BOTTOM 3 ĐIỆN LỰC', level=1)
        for _, row in bottom_3.iterrows():
            doc.add_paragraph(
                f"- {row['Điện lực']}: {row['Tổng công tơ']:,}/{row['Kế hoạch']:,} ({row['Tỷ lệ']:.2f}%)"
            )
        doc.add_heading('IV. KẾT LUẬN', level=1)
        doc.add_paragraph(
            "Với tốc độ hiện tại, toàn đơn vị sẽ chưa đạt kế hoạch vào 30/09/2025 nếu không tăng tốc."
        )

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    word_file = generate_docx()
    excel_file = generate_excel()

    st.download_button("📄 Tải báo cáo Word", data=word_file, file_name="Bao_cao_phan_tich_dien_luc.docx")
    st.download_button("📊 Tải báo cáo Excel", data=excel_file, file_name="Bao_cao_phan_tich_dien_luc.xlsx")
